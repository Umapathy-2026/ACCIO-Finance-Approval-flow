#!/usr/bin/env python3
"""Generate comprehensive Word documentation for Project ACCIO."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ── TITLE PAGE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJECT ACCIO')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(1, 105, 111)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Finance AR Ticketing System')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(108, 106, 104)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\nTechnical Documentation')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'\n\nDate: {datetime.now().strftime("%d %B %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(108, 106, 104)

doc.add_page_break()

# ── TABLE OF CONTENTS ──
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Project Overview',
    '2. Technical Stack & Dependencies',
    '3. Project Folder Structure',
    '4. Full Source Code',
    '    4.1 Root Files',
    '    4.2 Application Factory (app/__init__.py)',
    '    4.3 Database Models (app/models.py)',
    '    4.4 Routes',
    '    4.5 Utilities',
    '    4.6 Templates',
    '5. Recent UI Changes (Logo Integration)',
    '6. Production-Grade Enhancements (v2)',
    '7. v2.1 Security & Bug Fixes Audit',
    '8. Azure Deployment Guide',
    '9. User Guide',
    '10. Complete Task Checklist',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.style = doc.styles['Normal']

doc.add_page_break()

# ── 1. PROJECT OVERVIEW ──
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'ACCIO is a Finance Accounts Receivable (AR) Ticketing Web Application built with Python Flask. '
    'It enables users to submit financial issue tickets (e.g., billing discrepancies, logistics costs, '
    'price adjustments), which are routed to approvers for review and action. Admins manage users, roles, '
    'and the 21 predefined issue form types.'
)
doc.add_paragraph('Key features:')
for f in [
    'Role-based access: User (Requester), Approver, Admin',
    '21 configurable issue types with dynamic custom fields',
    'Full ticket lifecycle: Submit, Approve, Reject, Send Back for Clarification',
    'Email notifications via Office365 SMTP',
    'File attachments (PDF, DOC, XLS, images up to 10MB)',
    'Modern SaaS-style UI with Tailwind CSS + Inter font + Lucide icons',
    'SQLite database (swappable to Azure SQL)',
    'Swappable auth module (local email/password → MSAL/Entra ID)',
    'Swappable file storage (local → Azure Blob Storage)',
    'Vanilla JS toast notification system',
    'Client-side search, filter, sort and pagination',
    'Session timeout with 5-minute warning countdown',
    'Login lockout after 5 failed attempts',
    'Forgot password flow with secure reset tokens',
    'In-app notification bell with real-time polling',
    'Excel export (single ticket + bulk admin)',
    'Bulk approve/reject from approver queue',
    'Ticket reassignment between approvers',
    'Graceful user deactivation with ticket reassignment',
    'Unsaved changes warning, double-submit prevention',
]:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ── 2. TECHNICAL STACK ──
doc.add_heading('2. Technical Stack & Dependencies', level=1)
doc.add_heading('Backend', level=2)
table = doc.add_table(rows=9, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Package', 'Version', 'Purpose']):
    table.rows[0].cells[i].text = h
    for run in table.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
data = [
    ('Flask', '3.0.3', 'Web framework'),
    ('Flask-SQLAlchemy', '3.1.1', 'ORM for database'),
    ('Flask-Login', '0.6.3', 'Session management / auth'),
    ('Werkzeug', '3.0.3', 'WSGI toolkit / password hashing'),
    ('python-dotenv', '1.0.1', 'Environment variable loading'),
    ('openpyxl', '3.1.2', 'Excel export generation'),
    ('SQLite', 'Built-in', 'Local development database'),
    ('smtplib (stdlib)', '—', 'Email sending via Office365'),
]
for idx, (pkg, ver, purp) in enumerate(data, 1):
    table.rows[idx].cells[0].text = pkg
    table.rows[idx].cells[1].text = ver
    table.rows[idx].cells[2].text = purp

doc.add_heading('Frontend (CDN-loaded)', level=2)
table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(['Resource', 'Source', 'Purpose']):
    table2.rows[0].cells[i].text = h
    for run in table2.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
for idx, (r, s, p) in enumerate([
    ('Tailwind CSS', 'cdn.tailwindcss.com', 'Utility CSS framework'),
    ('Inter Font', 'Google Fonts', 'Sans-serif typeface'),
    ('Lucide Icons', 'unpkg.com/lucide@latest', 'Icon library'),
], 1):
    table2.rows[idx].cells[0].text = r
    table2.rows[idx].cells[1].text = s
    table2.rows[idx].cells[2].text = p

doc.add_page_break()

# ── 3. FOLDER STRUCTURE ──
doc.add_heading('3. Project Folder Structure', level=1)
structure = """\
accio/
├── .env                             # Environment variables (secrets)
├── app.py                           # Application entry point
├── requirements.txt                 # Python dependencies
├── generate_docs.py                 # (optional) Documentation generator
├── instance/
│   └── ticketing.db                 # SQLite database (auto-created)
└── app/
    ├── __init__.py                  # Flask factory, DB init, seed data
    ├── models.py                    # SQLAlchemy models (User, IssueForm, Ticket, ApprovalLog, Notification)
    ├── static/
    │   └── logo.png                 # Company logo
    ├── routes/
    │   ├── __init__.py              # (empty)
    │   ├── auth.py                  # Login/logout, lockout, forgot/reset password
    │   ├── requester.py             # Dashboard, create/view/clarify/export tickets, notifications page
    │   ├── approver.py              # Queue, approve/reject/send-back, reassign, bulk actions
    │   ├── admin.py                 # Dashboard, users, forms, export, unlock, deactivation
    │   └── api.py                   # Form fields AJAX, notification API
    ├── utils/
    │   ├── __init__.py              # (empty)
    │   ├── email.py                 # Email notification service
    │   └── storage.py               # File upload handler
    └── templates/
        ├── base.html                # App shell, design system, toast, session timeout, notifications, filters
        ├── auth/
        │   ├── login.html           # Split-screen login page
        │   ├── forgot_password.html # Email input for reset
        │   └── reset_password.html  # New password with validation
        ├── requester/
        │   ├── dashboard.html       # My Tickets + Needs Attention with search/filter/sort/pagination
        │   ├── new_ticket.html      # Create ticket form with dynamic fields
        │   ├── ticket_detail.html   # Full detail + timeline + clarify + export
        │   └── notifications.html   # Full notification list
        ├── approver/
        │   ├── queue.html           # Pending/Under Review with bulk checkboxes + filter/sort/pagination
        │   └── ticket_detail.html   # Detail + action buttons + reassign modal
        └── admin/
            ├── dashboard.html       # KPI cards + recent tickets + export button
            ├── users.html           # User list + create/edit/unlock/deactivate modals
            ├── forms.html           # Form list with toggle
            └── edit_form.html       # Field editor with add/remove/reorder"""

for line in structure.strip().split('\n'):
    p = doc.add_paragraph(line)
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)

doc.add_page_break()

# ── 4. FULL SOURCE CODE (section headers only - actual code from files) ──
doc.add_heading('4. Full Source Code', level=1)

def add_file_section(doc, filepath, heading_text):
    doc.add_heading(heading_text, level=2)
    filepath_full = os.path.join(os.path.dirname(__file__), filepath)
    try:
        with open(filepath_full, 'r', encoding='utf-8') as f:
            content = f.read()
    except:
        doc.add_paragraph(f'[File not found: {filepath_full}]')
        return
    p = doc.add_paragraph()
    run = p.add_run(f'File: {filepath}')
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    run.font.color.rgb = RGBColor(1, 105, 111)
    for line in content.split('\n'):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(7.5)

doc.add_heading('4.1 Root Files', level=2)
add_file_section(doc, '.env', '.env')
add_file_section(doc, 'requirements.txt', 'requirements.txt')
add_file_section(doc, 'app.py', 'app.py')
doc.add_page_break()

doc.add_heading('4.2 Application Factory', level=2)
add_file_section(doc, 'app/__init__.py', 'app/__init__.py')
doc.add_page_break()

doc.add_heading('4.3 Database Models', level=2)
add_file_section(doc, 'app/models.py', 'app/models.py')
doc.add_page_break()

doc.add_heading('4.4 Routes', level=2)
add_file_section(doc, 'app/routes/auth.py', 'auth.py — Authentication')
add_file_section(doc, 'app/routes/api.py', 'api.py — API Endpoints')
add_file_section(doc, 'app/routes/requester.py', 'requester.py — Requester Routes')
add_file_section(doc, 'app/routes/approver.py', 'approver.py — Approver Routes')
add_file_section(doc, 'app/routes/admin.py', 'admin.py — Admin Routes')
doc.add_page_break()

doc.add_heading('4.5 Utilities', level=2)
add_file_section(doc, 'app/utils/email.py', 'email.py — Email Notifications')
add_file_section(doc, 'app/utils/storage.py', 'storage.py — File Storage')
doc.add_page_break()

doc.add_heading('4.6 Templates', level=2)
add_file_section(doc, 'app/templates/base.html', 'base.html — App Shell & Design System')
add_file_section(doc, 'app/templates/auth/login.html', 'login.html')
add_file_section(doc, 'app/templates/auth/forgot_password.html', 'forgot_password.html')
add_file_section(doc, 'app/templates/auth/reset_password.html', 'reset_password.html')
add_file_section(doc, 'app/templates/requester/dashboard.html', 'requester/dashboard.html')
add_file_section(doc, 'app/templates/requester/new_ticket.html', 'requester/new_ticket.html')
add_file_section(doc, 'app/templates/requester/ticket_detail.html', 'requester/ticket_detail.html')
add_file_section(doc, 'app/templates/requester/notifications.html', 'requester/notifications.html')
add_file_section(doc, 'app/templates/approver/queue.html', 'approver/queue.html')
add_file_section(doc, 'app/templates/approver/ticket_detail.html', 'approver/ticket_detail.html')
add_file_section(doc, 'app/templates/admin/dashboard.html', 'admin/dashboard.html')
add_file_section(doc, 'app/templates/admin/users.html', 'admin/users.html')
add_file_section(doc, 'app/templates/admin/forms.html', 'admin/forms.html')
add_file_section(doc, 'app/templates/admin/edit_form.html', 'admin/edit_form.html')

doc.add_page_break()

# ── 5. RECENT UI CHANGES ──
doc.add_heading('5. Recent UI Changes — Logo Integration', level=1)
doc.add_heading('5.1 Logo File Placement', level=2)
doc.add_paragraph('The company logo was placed at app/static/logo.png for proper serving by Flask.')
doc.add_paragraph('File: app/static/logo.png (876KB, PNG format)', style='List Bullet')
doc.add_paragraph('URL: /static/logo.png (served automatically by Flask)', style='List Bullet')

doc.add_heading('5.2 Login Page — Left Panel Background', level=2)
doc.add_paragraph('The login page shows the ACCIO logo as a full-screen background image on the left split panel with a subtle teal overlay.')

doc.add_heading('5.3 App Header — Brand Visibility', level=2)
doc.add_paragraph('Logo image at 40px height with bold "ACCIO" brand text beside it in teal color.')

doc.add_page_break()

# ── 6. PRODUCTION-GRADE ENHANCEMENTS ──
doc.add_heading('6. Production-Grade Enhancements (v2)', level=1)

enhancements = [
    ('Toast Notifications', 'Vanilla JS toast system with 4 types, auto-dismiss, max 3 visible, stack vertically. Flask flash messages auto-converted to toasts.'),
    ('Search, Filter & Sort', 'Client-side filter bar on all ticket lists (My Tickets, Approver Queue, Admin). Debounced search, status dropdown, date range. Column sorting with sort arrows.'),
    ('Pagination', 'Client-side pagination with 10/20/50/100 per page. Page number buttons with ellipsis. "Showing X–Y of Z tickets" counter. Filters reset to page 1.'),
    ('Loading States & Double-Submit Prevention', 'submitWithSpinner() function disables buttons, replaces text with spinner + "Processing...", re-enables after 10s fallback. Unsaved changes warning via beforeunload.'),
    ('Session Timeout', '30-minute inactivity timer. Warning modal at 25 min with 5-minute countdown. "Stay Logged In" resets timer. Tracks mouse/key/touch events.'),
    ('Login Lockout', 'Tracks failed_attempts per email. Locks after 5 attempts for 30 minutes. Shows remaining attempts warning after 2nd failure. Admin unlock button on Users page.'),
    ('Forgot Password Flow', 'GET/POST /auth/forgot-password. Secure 48-byte reset token, expires 1 hour. HTML email with CTA button. Reset page with inline password validation (8 chars, 1 uppercase, 1 number).'),
    ('Ticket Number Copy Button', 'Copy button on every ticket number in tables and detail pages. Shows checkmark icon briefly on success.'),
    ('Audit Trail Export (Excel)', 'Single ticket export: GET /ticket/<id>/export (2 sheets: Ticket Details + Approval History). Admin bulk export: GET /admin/export with date/status/form filters. Styled Excel with teal headers and alternating rows.'),
    ('Bulk Approve/Reject', 'Checkbox column + Select All on approver queue. Sticky bottom action bar. POST /approver/bulk-action. Reject opens modal for common reason.'),
    ('Ticket Reassignment', '"Reassign" button on approver ticket detail. Modal with approver dropdown + reason. Logs reassignment, notifies both assignees and creator.'),
    ('Graceful User Deactivation', 'Checks for open tickets before deactivating. Shows interstitial modal with reassign dropdown. Bulk-reassigns tickets to selected approver before deactivation.'),
    ('In-App Notification Bell', 'Notification model with DB persistence. Bell icon with red badge, polled every 60s. Slide-down panel showing last 5 notifications. Mark all read, click to navigate. Full /notifications page.'),
]

for i, (title, desc) in enumerate(enhancements, 1):
    doc.add_heading(f'6.{i} {title}', level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ── 7. V2.1 SECURITY & BUG FIXES AUDIT ──
doc.add_heading('7. v2.1 Security & Bug Fixes Audit', level=1)
doc.add_paragraph(
    'The following section documents all security vulnerabilities, critical bugs, and logic issues '
    'that were identified and fixed during the v2.1 audit. This ensures the application meets '
    'enterprise security standards for a finance AR ticketing system.'
)

fixes = [
    ('SEC-01: CSRF Protection (Flask-WTF)',
     'Implemented CSRF protection using Flask-WTF CSRFProtect. Every form now includes a CSRF token, '
     'and API JSON endpoints are explicitly exempted. This prevents Cross-Site Request Forgery attacks '
     'where a malicious email or site could silently approve/reject tickets.'),
    ('SEC-02: HTTP Security Headers',
     'Added security response headers via Flask-Talisman: X-Frame-Options (DENY) prevents clickjacking, '
     'X-Content-Type-Options (nosniff) prevents MIME confusion, Content-Security-Policy restricts scripts/'
     'styles to trusted CDNs, and Strict-Transport-Security enforces HTTPS in production.'),
    ('SEC-03: Force Password Change on First Login',
     'The default admin account (admin@company.com / Admin@123) now has must_change_password=True set '
     'during seed. A @before_request handler enforces password reset for any user with this flag, '
     'preventing continued use of default credentials.'),
    ('SEC-04: SECRET_KEY Validation',
     'The application now validates SECRET_KEY at startup. In production mode, it raises RuntimeError '
     'if missing. In development, it shows a deprecation warning. This prevents silent fallback to a '
     'known public default key that would allow session forgery.'),
    ('SEC-06: HTML Injection Prevention in Emails',
     'All user-controlled content (ticket subject, comments, rejection reasons) passed to email templates '
     'is now escaped using markupsafe.escape(). This prevents HTML injection and phishing via styled '
     'malicious links in official ACCIO emails.'),
    ('SEC-07: Debug Mode Safeguard',
     'app.py now reads FLASK_DEBUG from environment instead of hardcoding debug=True. The Werkzeug '
     'debugger with its remote Python REPL is no longer accessible by running python app.py in staging/production.'),
    ('SEC-08: Rate Limiting',
     'Implemented Flask-Limiter with default limits of 500 requests/day and 100 requests/hour. '
     'Login endpoint is limited to 10 requests/minute, and forgot-password to 5 requests/hour. '
     'A custom error handler returns JSON for API routes and flash messages for web routes when rate limited.'),
    ('BUG-01: Broken Email "Review Ticket" URL',
     'Email notification functions now accept a review_url parameter containing the full absolute URL. '
     'Routes pass url_for(..., _external=True) to generate complete URLs. Previously, emails contained '
     'broken relative links like href="ACC-20260615-0001" or href="#".'),
    ('BUG-03: Deactivated User Session Persistence',
     'The @login_manager.user_loader now checks user.is_active and returns None (treated as unauthenticated '
     'by Flask-Login) when a user is deactivated. Previously, deactivated users could continue using the '
     'app until their session naturally expired (30 minutes).'),
    ('BUG-04: Bulk Action Email Notifications',
     'The bulk_action() route now calls send_ticket_approved() and send_ticket_rejected() with proper '
     'review URLs for each ticket. Previously, bulk approve/reject only created in-app notifications '
     'but never sent email alerts to requesters.'),
    ('BUG-05: "Sent to Fulfilment" Status Unreachable',
     'Ticket approval now automatically sets status to "Sent to Fulfilment" (SENT_TO_FULFILMENT) instead '
     'of "Approved". An additional ApprovalLog entry is created to record the fulfilment transition. '
     'This matches the email copy that tells users "approved and sent to fulfilment."'),
    ('BUG-06: Approver Ticket Orphaning on Deactivation',
     'The deactivation workflow now checks both tickets created BY the user AND tickets ASSIGNED TO the '
     'user. The deactivate modal provides two separate reassign dropdowns — one for created tickets and '
     'one for assigned tickets. Previously, only created tickets were checked, leaving assigned tickets orphaned.'),
    ('BUG-07: datetime.utcnow() Migration',
     'All occurrences of datetime.utcnow() throughout models.py, routes, and utils were replaced with '
     'datetime.now(timezone.utc). A utcnow() helper function is defined in models.py for DRY usage. '
     'This eliminates DeprecationWarning in Python 3.12+ and ensures future compatibility.'),
    ('LOGIC-01: REASSIGNED Action Added to Enum',
     'ApprovalAction enum now includes REASSIGNED = "Reassigned". Previously, reassignment was logged '
     'with a hardcoded string, making it inconsistent with other enum-based action tracking.'),
    ('LOGIC-02: updated_at Explicit on Status Changes',
     'The Ticket.updated_at column now defaults to utcnow() with onupdate=utcnow(). Additionally, '
     'every route that changes ticket status (approve, reject, send-back, reassign, bulk action) '
     'explicitly sets ticket.updated_at = datetime.now(timezone.utc). This ensures the field is '
     'always accurate even with bulk updates that bypass ORM events.'),
    ('ARCH-04: Admin Audit Log',
     'Added AdminAuditLog model with columns: performed_by, action, target_type, target_id, details (JSON), '
     'ip_address, and timestamp. Every admin action (user create/edit/unlock/deactivate, form toggle/edit, '
     'export trigger) is logged. A paginated audit log UI is available at /admin/audit-log.'),
    ('ARCH-05: Health Check Endpoint',
     'Added GET /api/health endpoint that verifies database connectivity by executing SELECT 1. '
     'Returns JSON {"status": "ok", "db": "connected"} on success or {"status": "error", "db": "<msg>"} '
     'with 503 status on failure. Required by Azure App Service for instance health monitoring.'),
    ('ARCH-07: Database Indexes',
     'Added composite indexes on tickets table: (assigned_to, current_status), (created_by, current_status), '
     '(created_at), and (form_id). Also added indexes on notification table: (user_id, is_read) and '
     '(created_at), and admin_audit_log table: (performed_by) and (timestamp).'),
    ('F-02: Admin "All Tickets" View',
     'Added /admin/tickets route with paginated, searchable ticket list for admins. Shows all tickets '
     'in the system with a 500-row cap, plus total count. Links to each ticket detail page.'),
    ('F-03: Approver Resolved Tickets Tab',
     'Approver queue now loads both active tickets (Pending, Under Review, Needs Clarification) and '
     'resolved tickets (Approved, Rejected, Sent to Fulfilment). Both lists are capped at 500 rows. '
     'Resolved tickets are sorted by updated_at descending.'),
]

for i, (title, desc) in enumerate(fixes, 1):
    doc.add_heading(f'7.{i} {title}', level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ── 8. AZURE DEPLOYMENT GUIDE ──
doc.add_heading('8. Azure Deployment Guide', level=1)
doc.add_paragraph('Changes required to deploy ACCIO on Microsoft Azure using minimal-cost services.')

doc.add_heading('7.1 Architecture (Minimal Cost)', level=2)
arch_data = [
    ('Azure App Service (B1)', '~$13/mo', 'Flask app hosting'),
    ('Azure SQL (Serverless)', '~$5/mo', 'Managed SQL database'),
    ('Azure Blob Storage (LRS)', '~$1/mo', 'File uploads'),
    ('Azure Communication Services', '~$0/mo', 'Email (pay per message)'),
    ('Microsoft Entra ID (Free)', '$0/mo', 'Authentication'),
]
table3 = doc.add_table(rows=len(arch_data)+1, cols=3)
table3.style = 'Light Grid Accent 1'
for i, h in enumerate(['Service', 'Est. Cost', 'Purpose']):
    table3.rows[0].cells[i].text = h
    for run in table3.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
for idx, (s, c, p) in enumerate(arch_data, 1):
    table3.rows[idx].cells[0].text = s
    table3.rows[idx].cells[1].text = c
    table3.rows[idx].cells[2].text = p

doc.add_paragraph('')
doc.add_paragraph('Estimated total: ~$19-25/month for low usage.', style='List Bullet')

doc.add_heading('7.2 File Changes for Azure', level=2)

changes = [
    ('Database', 'In app/__init__.py, change SQLALCHEMY_DATABASE_URI to Azure SQL connection string. Add pymssql to requirements.txt. Configure firewall rules.'),
    ('File Storage', 'Replace save_file() in app/utils/storage.py with Azure Blob Storage using azure-storage-blob. Create "uploads" container.'),
    ('Authentication', 'Replace app/routes/auth.py with MSAL. Register app in Azure Entra ID. Login redirects to Microsoft. Only auth.py needs changes.'),
    ('Email', 'Replace smtplib in app/utils/email.py with Azure Communication Services Email. Create Email Communication Service resource.'),
]

for title, desc in changes:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading('7.3 Deployment Steps', level=2)
steps = [
    'Create Web App (Linux, Python 3.13, B1 tier)',
    'Configure App Settings: SECRET_KEY, SQLALCHEMY_DATABASE_URI, etc.',
    'Enable "Always On"',
    'Deploy via GitHub Actions or ZIP deploy',
    'Startup command: gunicorn --bind=0.0.0.0:8000 app:app',
    'Add gunicorn to requirements.txt',
    'Configure SQLAlchemy pool settings',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_page_break()

# ── 9. USER GUIDE ──
doc.add_heading('9. User Guide', level=1)
doc.add_heading('8.1 Getting Started', level=2)
doc.add_paragraph('Navigate to http://127.0.0.1:5000 (local) or your deployed URL.')
doc.add_heading('8.2 Logging In', level=2)
doc.add_paragraph('Default: admin@company.com / Admin@123')
doc.add_heading('8.3 Creating a Ticket', level=2)
for s in ['Click "Create Ticket" in sidebar', 'Select issue type, fill dynamic fields', 'Add subject, description, attachment', 'Submit — approver notified via email']:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('8.4 Viewing & Clarifying Tickets', level=2)
doc.add_paragraph('My Tickets tab shows all tickets. Needs Attention tab shows tickets sent back for clarification. Use the search/filter bar to find specific tickets.')

doc.add_heading('8.5 Approving Tickets', level=2)
doc.add_paragraph('Approval Queue shows tickets requiring review. Use Approve/Send Back/Reject buttons. Bulk select checkbox for multi-action. Reassign button to transfer to another approver.')

doc.add_heading('8.6 Admin Functions', level=2)
doc.add_paragraph('Dashboard: KPI cards + recent tickets + Export to Excel button.')
doc.add_paragraph('Users: Create/edit/unlock/deactivate users. Graceful deactivation reassigns open tickets.')
doc.add_paragraph('Forms: Toggle active/inactive, edit custom fields.')

doc.add_heading('8.7 Notifications', level=2)
doc.add_paragraph('Bell icon in header shows unread count. Click to view last 5. Click "View all" for full list.')

doc.add_page_break()

# ── 10. TASK CHECKLIST ──
doc.add_heading('10. Complete Task Checklist', level=1)
doc.add_paragraph('This section lists every feature, page, and system component as individual trackable items. Use this as a project checklist.')

task_groups = [
    ('Auth System', [
        'ACCIO-1 Login page with email/password',
        'ACCIO-2 Login lockout (5 failed attempts → 30min lock)',
        'ACCIO-3 "Forgot your password?" link on login page',
        'ACCIO-4 Forgot password page (email input)',
        'ACCIO-5 Reset password token generation (48-byte, 1hr expiry)',
        'ACCIO-6 Reset password page with inline validation',
        'ACCIO-7 Password reset email template',
        'ACCIO-8 Logout route',
        'ACCIO-9 role_required decorator',
        'ACCIO-10 Role-based redirect after login',
    ]),
    ('User (Requester) Dashboard', [
        'ACCIO-11 My Tickets tab with table',
        'ACCIO-12 Needs Attention tab (Needs Clarification tickets)',
        'ACCIO-13 Search/filter bar (search, status dropdown, clear)',
        'ACCIO-14 Column sorting (asc/desc/reset)',
        'ACCIO-15 Pagination (10/20/50 per page)',
        'ACCIO-16 Empty state ("No tickets yet")',
        'ACCIO-17 Filtered empty state ("No tickets match your filters")',
        'ACCIO-18 Clickable rows → ticket detail',
        'ACCIO-19 Copy ticket number button on each row',
    ]),
    ('Create Ticket', [
        'ACCIO-20 Dropdown to select issue type (21 types)',
        'ACCIO-21 AJAX dynamic field loading via /api/form-fields/<id>',
        'ACCIO-22 Subject (required), Description (optional)',
        'ACCIO-23 File attachment (10MB limit)',
        'ACCIO-24 Submit button with spinner prevention',
        'ACCIO-25 Unsaved changes warning',
        'ACCIO-26 Auto-generate ticket number (ACC-YYYYMMDD-XXXX)',
        'ACCIO-27 Auto-assign approver (manager or fallback)',
        'ACCIO-28 ApprovalLog entry on submission',
        'ACCIO-29 Email notification to approver on creation',
    ]),
    ('Ticket Detail (Requester)', [
        'ACCIO-30 Full ticket information display',
        'ACCIO-31 Dynamic payload field values',
        'ACCIO-32 Activity timeline (ApprovalLog)',
        'ACCIO-33 Clarification form when status = Needs Clarification',
        'ACCIO-34 File upload with clarification',
        'ACCIO-35 Copy ticket number button',
        'ACCIO-36 Export to Excel button',
    ]),
    ('Approver Queue', [
        'ACCIO-37 Pending + Under Review tickets list',
        'ACCIO-38 Search/filter bar with search, status, date range',
        'ACCIO-39 Column sorting',
        'ACCIO-40 Pagination',
        'ACCIO-41 Bulk checkbox + Select All',
        'ACCIO-42 Sticky bulk action bar (Approve All / Reject All / Cancel)',
        'ACCIO-43 Copy ticket number button on each row',
    ]),
    ('Approver Actions', [
        'ACCIO-44 Approve with comment modal',
        'ACCIO-45 Reject with reason modal',
        'ACCIO-46 Send Back for Clarification modal',
        'ACCIO-47 Sticky action bar at bottom of ticket detail',
        'ACCIO-48 Reassign ticket to another approver',
        'ACCIO-49 Bulk approve (POST /approver/bulk-action)',
        'ACCIO-50 Bulk reject with common reason',
        'ACCIO-51 Email notifications on approve/reject/send-back',
        'ACCIO-52 In-app notification on approve/reject/clarify',
    ]),
    ('Admin Dashboard', [
        'ACCIO-53 KPI cards: Total, This Month, Pending, Approved',
        'ACCIO-54 Recent tickets table (last 10)',
        'ACCIO-55 Export to Excel button with date/status/form filter modal',
    ]),
    ('Admin User Management', [
        'ACCIO-56 List all users with role and manager',
        'ACCIO-57 Create user modal (name, email, password, role, approver)',
        'ACCIO-58 Edit user modal (role, approver, active, password)',
        'ACCIO-59 Unlock user button (show when locked)',
        'ACCIO-60 Deactivate user with grace (check open tickets)',
        'ACCIO-61 Deactivation interstitial modal with reassign dropdown',
    ]),
    ('Admin Form Management', [
        'ACCIO-62 List 21 issue forms with field count + active status',
        'ACCIO-63 Toggle active/inactive',
        'ACCIO-64 Edit form fields (add/remove/reorder)',
        'ACCIO-65 Field types: text, number, date, dropdown, email, file',
        'ACCIO-66 Required toggle per field',
        'ACCIO-67 Dropdown options editor (one per line)',
    ]),
    ('System-Wide Features', [
        'ACCIO-68 Toast notification system (4 types, auto-dismiss, stack)',
        'ACCIO-69 Session timeout (30 min, 5 min warning, countdown)',
        'ACCIO-70 Double-submit prevention (spinner + disable)',
        'ACCIO-71 Unsaved changes warning (beforeunload)',
        'ACCIO-72 Notification model + in-app bell',
        'ACCIO-73 Notification polling (every 60s)',
        'ACCIO-74 Notification panel (last 5) + full page',
        'ACCIO-75 Company logo on login + header',
    ]),
    ('Azure Deployment', [
        'ACCIO-76 Set up Azure App Service (B1, Linux, Python 3.13)',
        'ACCIO-77 Configure Azure SQL Database (Serverless)',
        'ACCIO-78 Configure Azure Blob Storage for file uploads',
        'ACCIO-79 Configure Azure Communication Services for email',
        'ACCIO-80 Configure Microsoft Entra ID for auth (MSAL)',
        'ACCIO-81 Update requirements.txt (add pymssql, azure-*)',
        'ACCIO-82 Update app/__init__.py for Azure SQL connection',
        'ACCIO-83 Update app/utils/storage.py for Azure Blob',
        'ACCIO-84 Update app/routes/auth.py for MSAL',
        'ACCIO-85 Update app/utils/email.py for ACS Email',
        'ACCIO-86 Set all App Settings in Azure Portal',
        'ACCIO-87 Deploy and test authentication flow',
        'ACCIO-88 Deploy and test ticket lifecycle',
        'ACCIO-89 Deploy and test email notifications',
        'ACCIO-90 Deploy and test file uploads',
    ]),
]

for group_name, task_items in task_groups:
    doc.add_heading(group_name, level=2)
    table_t = doc.add_table(rows=len(task_items)+1, cols=2)
    table_t.style = 'Light Grid Accent 1'
    table_t.rows[0].cells[0].text = 'Task ID'
    table_t.rows[0].cells[1].text = 'Description'
    for run in table_t.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True
    for run in table_t.rows[0].cells[1].paragraphs[0].runs:
        run.bold = True
    for idx, item in enumerate(task_items):
        parts = item.split(' ', 1)
        tid = parts[0]
        desc = parts[1] if len(parts) > 1 else ''
        table_t.rows[idx+1].cells[0].text = tid
        table_t.rows[idx+1].cells[1].text = desc

    doc.add_paragraph('')

# ── SAVE ──
output_path = os.path.join(os.path.dirname(__file__), 'ACCIO_Project_Documentation.docx')
doc.save(output_path)
print(f'Documentation saved to: {output_path}')